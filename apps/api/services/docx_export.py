"""
Markdown-to-DOCX converter using python-docx.

Converts markdown text to a professional Word document with proper
heading hierarchy, bold/italic inline formatting, and bulleted lists.
"""

from __future__ import annotations

import re
from html.parser import HTMLParser
from io import BytesIO
from typing import Optional

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class _DocxBuilder(HTMLParser):
    """Walks HTML elements and builds a python-docx Document."""

    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self._current_paragraph = None
        self._bold = False
        self._italic = False
        self._in_heading = 0  # 0 = not in heading, 1-6 = heading level
        self._heading_text = ""
        self._in_list = False
        self._in_code_block = False
        self._code_text = ""

    def _ensure_paragraph(self):
        if self._current_paragraph is None:
            self._current_paragraph = self.doc.add_paragraph()
        return self._current_paragraph

    def _add_run(self, text: str):
        if not text:
            return
        p = self._ensure_paragraph()
        run = p.add_run(text)
        run.bold = self._bold
        run.italic = self._italic

    def handle_starttag(self, tag: str, attrs: list):
        tag = tag.lower()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._in_heading = int(tag[1])
            self._heading_text = ""
            self._current_paragraph = None
            return

        if tag == "p":
            self._current_paragraph = None
            return

        if tag in ("strong", "b"):
            self._bold = True
            return

        if tag in ("em", "i"):
            self._italic = True
            return

        if tag == "li":
            self._current_paragraph = self.doc.add_paragraph(style="List Bullet")
            return

        if tag in ("ul", "ol"):
            self._in_list = True
            self._current_paragraph = None
            return

        if tag == "br":
            if self._current_paragraph:
                self._current_paragraph.add_run("\n")
            return

        if tag == "hr":
            self._current_paragraph = None
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            run.font.size = Pt(8)
            run.font.color.theme_color = None
            self._current_paragraph = None
            return

        if tag == "pre":
            self._in_code_block = True
            self._code_text = ""
            return

        if tag == "code" and not self._in_code_block:
            # Inline code — just use italic monospace-like rendering
            self._italic = True
            return

    def handle_endtag(self, tag: str):
        tag = tag.lower()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = min(self._in_heading, 4)  # python-docx supports levels 0-9
            self.doc.add_heading(self._heading_text.strip(), level=level)
            self._in_heading = 0
            self._heading_text = ""
            self._current_paragraph = None
            return

        if tag == "p":
            self._current_paragraph = None
            return

        if tag in ("strong", "b"):
            self._bold = False
            return

        if tag in ("em", "i"):
            self._italic = False
            return

        if tag in ("ul", "ol"):
            self._in_list = False
            self._current_paragraph = None
            return

        if tag == "li":
            self._current_paragraph = None
            return

        if tag == "pre":
            self._in_code_block = False
            if self._code_text.strip():
                p = self.doc.add_paragraph()
                run = p.add_run(self._code_text.strip())
                run.font.size = Pt(9)
                run.font.name = "Courier New"
            self._code_text = ""
            self._current_paragraph = None
            return

        if tag == "code" and not self._in_code_block:
            self._italic = False
            return

    def handle_data(self, data: str):
        if self._in_heading:
            self._heading_text += data
            return

        if self._in_code_block:
            self._code_text += data
            return

        # Skip pure whitespace between block elements
        if not data.strip() and self._current_paragraph is None:
            return

        self._add_run(data)


def markdown_to_docx(text: str, title: str = "Document") -> BytesIO:
    """
    Convert markdown text to a DOCX file.

    Args:
        text: Markdown or plain text content.
        title: Document title (shown in header).

    Returns:
        BytesIO buffer containing the DOCX file.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Convert markdown to HTML
    html = markdown.markdown(
        text,
        extensions=["extra", "sane_lists"],
    )

    # Parse HTML into DOCX elements
    builder = _DocxBuilder(doc)
    builder.feed(html)

    # Post-process: expand fill-in blanks and signature lines
    _expand_fill_ins(doc)

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _expand_fill_ins(doc: Document):
    """
    Post-process the document to add space for fill-in fields and signatures.

    - Underscores (___) → wide fill-in line (tab stops)
    - "Signature" lines → 2 blank lines above the underline for signing space
    """
    from docx.shared import Pt as _Pt
    from docx.oxml.ns import qn

    for paragraph in doc.paragraphs:
        full_text = paragraph.text

        # Detect signature blocks: lines containing "Signature" or "Sign:" with underscores
        is_signature_line = bool(
            re.search(r"(?i)(signature|sign\s*here|signed|witness)", full_text)
            and re.search(r"_{3,}", full_text)
        )

        if is_signature_line:
            # Add 2 blank lines before signature for signing space
            p_element = paragraph._element
            parent = p_element.getparent()
            idx = list(parent).index(p_element)

            for _ in range(2):
                blank = doc.add_paragraph()
                blank.paragraph_format.space_before = _Pt(0)
                blank.paragraph_format.space_after = _Pt(0)
                # Move the blank paragraph before the signature line
                parent.remove(blank._element)
                parent.insert(idx, blank._element)

        # Expand short underscores to wider fill-in lines
        for run in paragraph.runs:
            if re.search(r"_{3,}", run.text):
                # Replace short underscores with longer ones for fill-in space
                run.text = re.sub(r"_{3,}", "____________________________", run.text)
